from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
IN = ROOT / 'docs' / 'modules.docx'
OUT = ROOT / 'docs' / 'modules-edited.docx'
IMAGE = ROOT / 'tools' / 'vb-word-sample' / 'sample-image.png'

if not IN.exists():
    print('Input not found:', IN)
    raise SystemExit(1)

doc = Document(str(IN))

# Set core properties
try:
    doc.core_properties.title = 'Enterprise Console — Modules'
    doc.core_properties.author = 'Dev Team'
except Exception:
    pass

# Ensure Normal style font
if 'Normal' in doc.styles:
    try:
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
    except Exception:
        pass

# Create a cover page at the start (insert a new section)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

cover_section = doc.add_section()
# Move to cover section header/footer
para = cover_section.header.paragraphs[0]
para.text = ''

# Cover title centered
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
trun = title.add_run('Enterprise Console')
trun.bold = True
trun.font.size = Pt(28)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.add_run('Modules & Sub-functionalities')
subtitle_run.italic = True
subtitle_run.font.size = Pt(14)

meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run(f'Generated on {__import__("datetime").datetime.now().date().isoformat()}')

# Insert cover image if available (centered)
if IMAGE.exists():
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    try:
        doc.add_picture(str(IMAGE), width=Inches(4))
    except Exception:
        # fallback: ignore image errors
        pass

# Add a page break after cover
doc.add_page_break()

# Insert a TOC placeholder paragraph with instruction
toc_para = doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
if 'Intense Quote' in doc.styles:
    toc_para.style = doc.styles['Intense Quote']

# Insert a simple TOC field using raw XML so Word can update it later
def add_simple_toc(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fld)

add_simple_toc(toc_para)

# Small replacement in paragraphs to mark edited status (non-destructive)
for p in doc.paragraphs:
    if 'Enterprise Console' in p.text and ' (edited)' not in p.text:
        p.text = p.text.replace('Enterprise Console', 'Enterprise Console (edited)')

# Add a small 2x2 table example at the end for demonstration
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for r in range(2):
    for c in range(2):
        table.cell(r, c).text = f'R{r+1}C{c+1}'

# Insert image again at the end if exists
if IMAGE.exists():
    try:
        doc.add_picture(str(IMAGE), width=Inches(2))
    except Exception:
        pass

# Add a merged-cell example table (2 rows, 3 cols, merge top row across 3 cols)
mt = doc.add_table(rows=2, cols=3)
mt.style = 'Table Grid'
mt.rows[0].cells[0].merge(mt.rows[0].cells[1]).merge(mt.rows[0].cells[2])
mt.rows[0].cells[0].text = 'Merged header across 3 columns'
mt.rows[1].cells[0].text = 'Cell A1'
mt.rows[1].cells[1].text = 'Cell A2'
mt.rows[1].cells[2].text = 'Cell A3'

# Header/footer example (simple)
section = doc.sections[0]
section.header.is_linked_to_previous = False
header_para = section.header.paragraphs[0]
header_para.text = 'Enterprise Console — Draft'
try:
    header_para.style.font.size = Pt(9)
except Exception:
    pass

footer_para = section.footer.paragraphs[0]
footer_para.text = f'Edited: {__import__("datetime").datetime.now().isoformat()}'
try:
    footer_para.style.font.size = Pt(9)
except Exception:
    pass

# Save a regenerated modules.docx (overwrites original) so CI can keep canonical copy if desired
regenerated = ROOT / 'docs' / 'modules.docx'
doc.save(str(OUT))
doc.save(str(regenerated))
print('Saved:', OUT, 'and regenerated canonical:', regenerated)
 
