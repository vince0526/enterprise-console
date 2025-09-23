from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / 'docs' / 'modules.md'
OUT = ROOT / 'docs' / 'modules-outline-numbered-level5-new.docx'

if not MD.exists():
    print('Markdown source not found:', MD)
    raise SystemExit(1)

lines = MD.read_text(encoding='utf-8').splitlines()

doc = Document()
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

# Cover
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('Enterprise Console — Modules').bold = True

doc.add_paragraph('Generated on: ' + __import__('datetime').datetime.now().isoformat())
doc.add_page_break()

# TOC placeholder
doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
doc.add_page_break()

# counters up to 5 levels
c = [0,0,0,0,0]

for line in lines:
    s = line.strip()
    if not s:
        continue
    if s.startswith('# '):
        c[0] += 1
        c[1]=c[2]=c[3]=c[4]=0
        num = f"{c[0]}"
        doc.add_heading(f"{num} {s[2:].strip()}", level=1)
    elif s.startswith('## '):
        c[1] += 1
        c[2]=c[3]=c[4]=0
        num = f"{c[0]}.{c[1]}"
        doc.add_heading(f"{num} {s[3:].strip()}", level=2)
    elif s.startswith('### '):
        c[2] += 1
        c[3]=c[4]=0
        num = f"{c[0]}.{c[1]}.{c[2]}"
        doc.add_heading(f"{num} {s[4:].strip()}", level=3)
    elif s.startswith('#### '):
        c[3] += 1
        c[4]=0
        num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}"
        doc.add_heading(f"{num} {s[5:].strip()}", level=4)
    elif s.startswith('##### '):
        c[4] += 1
        num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}.{c[4]}"
        doc.add_heading(f"{num} {s[6:].strip()}", level=5)
    elif s.startswith('- '):
        text = s[2:].strip()
        if ':' in text and ('resources/' in text or 'app/' in text or 'routes/' in text):
            key,val = text.split(':',1)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(key.strip()+': ').bold=True
            p.add_run(val.strip())
        else:
            doc.add_paragraph(text, style='List Bullet')
    else:
        doc.add_paragraph(s)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print('Saved level-5 numbered DOCX (new):', OUT)
