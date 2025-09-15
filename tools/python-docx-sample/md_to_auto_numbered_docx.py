from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# This script uses a simple approach: write headings as normal, then apply a
# numbered style by setting the style name to one that uses multilevel numbering
# if available. python-docx doesn't provide high-level API for numbering styles,
# so we create heading paragraphs and set their style to 'Heading 1','Heading 2'
# etc. Word will display numbering if the template/normal.dotx has multilevel
# numbering enabled or the user applies numbering in Word. For more robust
# programmatic numbering you'd need to inject numbering definitions via xml.

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / 'docs' / 'modules.md'
OUT = ROOT / 'docs' / 'modules-outline-auto-numbered.docx'

if not MD.exists():
    print('Markdown source not found:', MD)
    raise SystemExit(1)

lines = MD.read_text(encoding='utf-8').splitlines()

doc = Document()
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

# Cover
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run('Enterprise Console — Modules')
r.bold = True
r.font.size = Pt(24)

doc.add_paragraph('Generated on: ' + __import__('datetime').datetime.now().isoformat())
doc.add_page_break()

# TOC placeholder
doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
doc.add_page_break()

# counters (not used for numbering; Word will add numbers visually if multilevel is set)
for line in lines:
    s = line.strip()
    if not s:
        continue
    if s.startswith('# '):
        doc.add_paragraph(s[2:].strip(), style='Heading 1')
    elif s.startswith('## '):
        doc.add_paragraph(s[3:].strip(), style='Heading 2')
    elif s.startswith('### '):
        doc.add_paragraph(s[4:].strip(), style='Heading 3')
    elif s.startswith('- '):
        text = s[2:].strip()
        # Bold key if path-like
        if ':' in text and ('resources/' in text or 'app/' in text or 'routes/' in text):
            key, val = text.split(':', 1)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(key.strip() + ': ').bold = True
            p.add_run(val.strip())
        else:
            doc.add_paragraph(text, style='List Bullet')
    else:
        doc.add_paragraph(s)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print('Saved:', OUT)
