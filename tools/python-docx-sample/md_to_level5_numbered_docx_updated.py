from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / 'docs' / 'modules.md'
OUT = ROOT / 'docs' / 'modules-outline-numbered-level5-updated-new.docx'

if not MD.exists():
    print('Markdown source not found:', MD)
    raise SystemExit(1)

lines = MD.read_text(encoding='utf-8').splitlines()

doc = Document()
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

# Cover
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('Enterprise Console — Modules').bold = True

doc.add_paragraph('Generated on: ' + __import__('datetime').datetime.now().isoformat())
doc.add_page_break()

# TOC placeholder
doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
doc.add_page_break()

# counters up to 5 levels
c = [0,0,0,0,0]

def strip_leading_numbers(text):
    # remove patterns like '1.2.3 ' or '1. ' at start
    return re.sub(r'^\s*\d+(?:\.\d+)*\s*', '', text)

for i, line in enumerate(lines):
    s = line.rstrip()
    if not s.strip():
        continue
    stripped = s.lstrip()
    indent = len(s) - len(stripped)
    # headings
    if stripped.startswith('# '):
        clean = strip_leading_numbers(stripped[2:].strip())
        c[0] += 1
        c[1]=c[2]=c[3]=c[4]=0
        num = f"{c[0]}"
        doc.add_heading(f"{num} {clean}", level=1)
    elif stripped.startswith('## '):
        clean = strip_leading_numbers(stripped[3:].strip())
        c[1] += 1
        c[2]=c[3]=c[4]=0
        num = f"{c[0]}.{c[1]}"
        doc.add_heading(f"{num} {clean}", level=2)
    elif stripped.startswith('### '):
        clean = strip_leading_numbers(stripped[4:].strip())
        c[2] += 1
        c[3]=c[4]=0
        num = f"{c[0]}.{c[1]}.{c[2]}"
        doc.add_heading(f"{num} {clean}", level=3)
    elif stripped.startswith('#### '):
        clean = strip_leading_numbers(stripped[5:].strip())
        c[3] += 1
        c[4]=0
        num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}"
        doc.add_heading(f"{num} {clean}", level=4)
    elif stripped.startswith('##### '):
        clean = strip_leading_numbers(stripped[6:].strip())
        c[4] += 1
        num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}.{c[4]}"
        doc.add_heading(f"{num} {clean}", level=5)
    elif stripped.startswith('- '):
        # If list item appears under a level-3 heading and is indented, promote it to level-4 heading
        text = stripped[2:].strip()
        # look back for last heading levels to determine context
        prev = None
        # find previous non-empty heading line in reverse
        for j in range(i-1, -1, -1):
            pl = lines[j].strip()
            if not pl:
                continue
            if pl.startswith('### '):
                prev = 3
                break
            if pl.startswith('#### '):
                prev = 4
                break
            if pl.startswith('## '):
                prev = 2
                break
            if pl.startswith('# '):
                prev = 1
                break
        if prev == 3:
            # promote to heading level 4
            c[3] += 1
            c[4]=0
            num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}"
            clean = strip_leading_numbers(text)
            doc.add_heading(f"{num} {clean}", level=4)
        elif prev == 4:
            # promote to heading level 5
            c[4] += 1
            num = f"{c[0]}.{c[1]}.{c[2]}.{c[3]}.{c[4]}"
            clean = strip_leading_numbers(text)
            doc.add_heading(f"{num} {clean}", level=5)
        else:
            # normal bullet
            if ':' in text and ('resources/' in text or 'app/' in text or 'routes/' in text):
                key,val = text.split(':',1)
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(key.strip()+': ').bold=True
                p.add_run(val.strip())
            else:
                doc.add_paragraph(text, style='List Bullet')
    else:
        doc.add_paragraph(stripped)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print('Saved level-5 numbered DOCX (updated):', OUT)
