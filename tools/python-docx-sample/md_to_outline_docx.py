import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / 'docs' / 'modules.md'
OUT = ROOT / 'docs' / 'modules-outline-structured.docx'

if len(sys.argv) > 1:
    MD = Path(sys.argv[1])
if len(sys.argv) > 2:
    OUT = Path(sys.argv[2])

if not MD.exists():
    print('Markdown source not found:', MD)
    raise SystemExit(1)

text = MD.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()

# Styles: Normal font
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

# Cover
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run('Enterprise Console — Modules')
title_run.bold = True
title_run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.add_run('Developer Reference — Modules & Sub-functionalities')
subtitle_run.italic = True
subtitle_run.font.size = Pt(12)

doc.add_paragraph('Generated on: ' + __import__('datetime').datetime.now().isoformat())
doc.add_page_break()

# TOC placeholder
toc_para = doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
toc_para._p.append(fld)
doc.add_page_break()

# Numbering state: counts per level
num_stack = []  # e.g. [1, 2] -> 1.2

def push_level(level):
    # level is 1,2,3 -> ensure length
    while len(num_stack) < level:
        num_stack.append(0)

def bump_at(level):
    push_level(level)
    num_stack[level-1] += 1
    # reset deeper levels
    for i in range(level, len(num_stack)):
        num_stack[i] = 0
    # return string like '1.2.3'
    return '.'.join(str(x) for x in num_stack if x > 0)

for line in lines:
    s = line.rstrip()
    if not s.strip():
        continue
    s = s.lstrip()
    if s.startswith('# '):
        num = bump_at(1)
        h = doc.add_heading(f'{num} {s[2:].strip()}', level=1)
    elif s.startswith('## '):
        num = bump_at(2)
        h = doc.add_heading(f'{num} {s[3:].strip()}', level=2)
    elif s.startswith('### '):
        num = bump_at(3)
        h = doc.add_heading(f'{num} {s[4:].strip()}', level=3)
    elif s.startswith('- '):
        text = s[2:].strip()
        p = doc.add_paragraph()
        if ':' in text and ('resources/' in text or 'app/' in text or 'routes/' in text):
            parts = text.split(':', 1)
            p.add_run(parts[0].strip() + ': ').bold = True
            p.add_run(parts[1].strip())
            p.style = 'List Bullet'
        else:
            p.add_run(text)
            p.style = 'List Bullet'
    elif s.startswith('* '):
        p = doc.add_paragraph(s[2:].strip(), style='List Bullet')
    elif s[0].isdigit() and s[1:3] == ') ':
        p = doc.add_paragraph(s[3:].strip(), style='List Number')
    else:
        p = doc.add_paragraph(s)

# Save as numbered outline
OUT = ROOT / 'docs' / 'modules-outline-numbered.docx'
doc.save(str(OUT))
print('Saved numbered outline DOCX:', OUT)
