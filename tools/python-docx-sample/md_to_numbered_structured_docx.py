from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / 'docs' / 'modules.md'
OUT = ROOT / 'docs' / 'modules-outline-numbered-structured.docx'

if not MD.exists():
    print('Markdown source not found:', MD)
    raise SystemExit(1)

lines = MD.read_text(encoding='utf-8').splitlines()

doc = Document()
# Normal
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

# Cover
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run('Enterprise Console — Modules')
title_run.bold = True
title_run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.add_run('Developer Reference — Modules & Sub-functionalities')
subtitle_run.italic = True
subtitle_run.font.size = Pt(12)

doc.add_paragraph('Generated on: ' + __import__('datetime').datetime.now().isoformat())
doc.add_page_break()

# TOC
toc_para = doc.add_paragraph('Table of Contents (update fields in Word: Select All -> F9)')
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
toc_para._p.append(fld)
doc.add_page_break()

# Numbering counters
c1 = 0
c2 = 0
c3 = 0

for line in lines:
    s = line.strip()
    if not s:
        continue
    if s.startswith('# '):
        c1 += 1
        c2 = 0
        c3 = 0
        heading = f"{c1}. {s[2:].strip()}"
        doc.add_heading(heading, level=1)
    elif s.startswith('## '):
        c2 += 1
        c3 = 0
        heading = f"{c1}.{c2} {s[3:].strip()}"
        doc.add_heading(heading, level=2)
    elif s.startswith('### '):
        c3 += 1
        heading = f"{c1}.{c2}.{c3} {s[4:].strip()}"
        doc.add_heading(heading, level=3)
    elif s.startswith('- '):
        text = s[2:].strip()
        # Format items: detect 'Key: value' and bold Key
        if ':' in text and ('resources/' in text or 'app/' in text or 'routes/' in text):
            key, val = text.split(':', 1)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(key.strip() + ': ').bold = True
            p.add_run(val.strip())
        else:
            p = doc.add_paragraph(text, style='List Bullet')
    elif s.startswith('* '):
        doc.add_paragraph(s[2:].strip(), style='List Bullet')
    elif s[0].isdigit() and s[1:3] == ') ':
        doc.add_paragraph(s[3:].strip(), style='List Number')
    else:
        doc.add_paragraph(s)

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print('Saved:', OUT)
