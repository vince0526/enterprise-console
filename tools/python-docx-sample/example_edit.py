from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / 'docs' / 'modules.docx'
OUTPUT = ROOT / 'docs' / 'modules.edited-by-script.docx'

if not INPUT.exists():
    print('Source file not found:', INPUT)
    raise SystemExit(1)

print('Reading:', INPUT)
doc = Document(str(INPUT))

# Append an edit note
p = doc.add_paragraph()
p.add_run('Edited by example_edit.py').italic = True

# Save
doc.save(str(OUTPUT))
print('Saved new file:', OUTPUT)
