#!/usr/bin/env python3
"""
Simple Markdown to DOCX converter for the project's `docs/modules.md`.
- Preserves headings (levels 1-3) and paragraphs.
- Requires: python-docx

Usage: python scripts/md_to_docx.py
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt
except Exception as e:
    print("Missing dependency 'python-docx'. Install with: pip install python-docx")
    raise

ROOT = Path(__file__).resolve().parents[1]
MD_FILE = ROOT / 'docs' / 'modules.md'
OUT_FILE = ROOT / 'docs' / 'modules.docx'

if not MD_FILE.exists():
    print(f"Markdown file not found: {MD_FILE}")
    sys.exit(1)

text = MD_FILE.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
styles = doc.styles

# Optional: tweak Normal font size
if 'Normal' in styles:
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(11)

for line in lines:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph('')
        continue
    # headings
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:].strip(), level=1)
        continue
    if stripped.startswith('## '):
        doc.add_heading(stripped[3:].strip(), level=2)
        continue
    if stripped.startswith('### '):
        doc.add_heading(stripped[4:].strip(), level=3)
        continue
    # bullet lists (simple)
    if stripped.startswith('- '):
        p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        continue
    if stripped.startswith('  - '):
        p = doc.add_paragraph(stripped[4:].strip(), style='List Bullet 2')
        continue
    # numbered lists
    if stripped[0:2].isdigit() and stripped[2:4] == ') ':
        doc.add_paragraph(stripped[4:].strip(), style='List Number')
        continue

    # default paragraph
    doc.add_paragraph(stripped)

# Save
OUT_FILE.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_FILE))
print(f"Written: {OUT_FILE}")
