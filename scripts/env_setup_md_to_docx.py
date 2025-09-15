#!/usr/bin/env python3
"""Convert docs/environment-setup.md to docs/environment-setup.docx.
Requires: python-docx
Usage: python scripts/env_setup_md_to_docx.py
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    print("Missing dependency 'python-docx'. Install with: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
MD_FILE = ROOT / 'docs' / 'environment-setup.md'
OUT_FILE = ROOT / 'docs' / 'environment-setup.docx'

if not MD_FILE.exists():
    print(f"Markdown file not found: {MD_FILE}")
    sys.exit(1)

text = MD_FILE.read_text(encoding='utf-8')
lines = text.splitlines()

def flush_table(pending_rows, document):
    if not pending_rows:
        return
    # Remove separator (---) row if present as second line
    rows_clean = []
    for r in pending_rows:
        rows_clean.append(r)
    # Interpret first row as header
    header = rows_clean[0]
    body = rows_clean[1:]
    table = document.add_table(rows=len(rows_clean), cols=len(header))
    table.style = 'Light List Accent 1' if 'Light List Accent 1' in [s.name for s in document.styles] else table.style
    # Header
    for ci, cell_text in enumerate(header):
        table.rows[0].cells[ci].text = cell_text
    # Body
    for ri, row_vals in enumerate(body, start=1):
        for ci, cell_text in enumerate(row_vals):
            table.rows[ri].cells[ci].text = cell_text

pending_table = []
in_table = False

doc = Document()
# Basic normal style tweak
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
except Exception:
    pass

for line in lines:
    raw = line.rstrip('\n')
    stripped = raw.strip()
    if not stripped:
        # Flush any active table
        if in_table:
            flush_table(pending_table, doc)
            pending_table = []
            in_table = False
        doc.add_paragraph('')
        continue
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:].strip(), level=1); continue
    if stripped.startswith('## '):
        doc.add_heading(stripped[3:].strip(), level=2); continue
    if stripped.startswith('### '):
        doc.add_heading(stripped[4:].strip(), level=3); continue
    if stripped.startswith('- '):
        doc.add_paragraph(stripped[2:].strip(), style='List Bullet'); continue
    if stripped.startswith('  - '):
        doc.add_paragraph(stripped[4:].strip(), style='List Bullet 2'); continue
    # Ordered list pattern "1) text" or "1. text"
    if (len(stripped) > 3 and stripped[0].isdigit() and (stripped[1:3] == ') ' or stripped[1] == '.' )):
        doc.add_paragraph(stripped[stripped.find(' ')+1:].strip(), style='List Number'); continue
    # Table detection: line starting and ending with | with at least one more |
    if stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2:
        # Parse row
        parts = [c.strip() for c in stripped.split('|')[1:-1]]
        # Detect separator row (---) and ignore except for structure
        if all(p.replace('-', '').replace(':', '') == '' for p in parts):
            # Table separator, ignore storing but ensure table mode
            in_table = True
            continue
        pending_table.append(parts)
        in_table = True
        continue
    else:
        if in_table:
            flush_table(pending_table, doc)
            pending_table = []
            in_table = False
    doc.add_paragraph(stripped)

if in_table and pending_table:
    flush_table(pending_table, doc)

OUT_FILE.parent.mkdir(exist_ok=True, parents=True)
doc.save(OUT_FILE)
print(f"Written: {OUT_FILE}")
